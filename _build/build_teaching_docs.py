from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "Teaching_Package"
BUILD_DIR = OUT_DIR / "_build"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F8FC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE6"
TEXT_GRAY = "4B5563"
GREEN = "E9F5EC"
GOLD = "FFF4CC"
RED = "FDECEC"
WHITE = "FFFFFF"
BLACK = "000000"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_borders(table, color: str = MID_GRAY, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Column widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_keep_with_next(paragraph, keep: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = keep


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=TEXT_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [style.name for style in styles]:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(NAVY)
    code_style.paragraph_format.left_indent = Inches(0.15)
    code_style.paragraph_format.right_indent = Inches(0.15)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.text = running_label
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        set_run_font(run, size=8.5, color=TEXT_GRAY, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)


def add_cover(doc: Document, kicker: str, title: str, subtitle: str,
              metadata: Sequence[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(kicker.upper())
    set_run_font(r, size=10.5, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, size=29, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(subtitle)
    set_run_font(r, size=14, color=TEXT_GRAY)

    table = doc.add_table(rows=len(metadata), cols=2)
    set_fixed_table_geometry(table, [2400, 6960])
    set_table_borders(table, color=WHITE, size=0)
    for i, (label, value) in enumerate(metadata):
        shade_cell(table.cell(i, 0), LIGHT_BLUE)
        shade_cell(table.cell(i, 1), PALE_BLUE)
        p1 = table.cell(i, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=10, color=NAVY, bold=True)
        p2 = table.cell(i, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Evidence before elegance. A chart is not finished until it reconciles.")
    set_run_font(r, size=13, color=DARK_BLUE, bold=True)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill: str = PALE_BLUE,
                accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=fill, size=1)
    shade_cell(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: Iterable[str], style: str = "List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    add_bullets(doc, items, style="List Number")


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths_dxa: Sequence[int], header_fill: str = LIGHT_BLUE,
              first_col_bold: bool = False, font_size: float = 9.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_fixed_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for i, row_values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for j, value in enumerate(row_values):
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK, bold=(first_col_bold and j == 0))
            if i % 2 == 0:
                shade_cell(cells[j], "FAFBFC")
    set_fixed_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.keep_together = True
        p.add_run(line if line else " ")


def add_section_start(doc: Document, title: str, lead: str | None = None) -> None:
    doc.add_heading(title, level=1)
    if lead:
        p = doc.add_paragraph()
        r = p.add_run(lead)
        set_run_font(r, size=11.5, color=NAVY, bold=True)


def add_image(doc: Document, image_path: Path, width_inches: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    set_run_font(r, size=8.5, color=TEXT_GRAY, italic=True)


def build_student_assignment() -> Path:
    doc = Document()
    style_document(doc, "FINC 310 | P&G Financial Story | Student Assignment")
    add_cover(
        doc,
        "FINC 310 Student Group Assignment",
        "Procter & Gamble Financial Story",
        "Use audited 10-K data, AI, and Python to explain profit, cash flow, growth, and return on equity.",
        [
            ("Company", "The Procter & Gamble Company (NYSE: PG)"),
            ("Team", "Six roles; combine roles if your team has fewer than six members"),
            ("Data", "Fiscal year ended June 30, 2025; USD millions except EPS"),
            ("Outputs", "Four visualizations, one dashboard, code, AI audit, interpretation, and presentation"),
            ("Standard", "AI use is required, documented, and independently verified"),
        ],
    )

    add_section_start(
        doc,
        "Quick Start",
        "Your team will build one financial story through four analytical lenses and prove that every result is reproducible.",
    )
    add_numbered(doc, [
        "Open PG_FinancialStory_Student.xlsx and read the 00_ReadMe and 01_Input worksheets.",
        "Assign the six team roles. Every member remains accountable for the entire submission.",
        "Complete the manual formulas before asking AI to generate code.",
        "Use the required prompt and REM structure to create one executable Python program.",
        "Run the six accounting and finance checks before accepting any chart.",
        "Interpret the four visuals as one integrated financial story, not four disconnected pictures.",
        "Submit the complete package and deliver a five-minute team presentation.",
    ])
    add_callout(
        doc,
        "Non-negotiable standard",
        "A polished chart that does not reconcile to audited data is unacceptable. Correct code that no team member can explain is also unacceptable.",
        fill=GOLD,
        accent="7A5A00",
    )

    add_section_start(doc, "1. Project Overview")
    doc.add_paragraph(
        "P&G is a public company whose financial statements connect operating decisions to cash flow, financing, risk, and shareholder value. Your team will use the supplied standardized workbook and audited annual report to create four complementary visualizations: an income statement Sankey, a free cash flow waterfall, a five-year revenue and EPS trend, and a DuPont ROE decomposition."
    )
    doc.add_paragraph(
        "The four visuals must tell one story. The Sankey explains where sales dollars went. The waterfall explains how operating profit becomes several definitions of free cash flow. The trend shows how the business changed over five fiscal years. DuPont explains whether ROE comes from profitability, asset use, or leverage."
    )
    add_table(
        doc,
        ["Question", "Evidence your team must produce"],
        [
            ("How profitable is the business?", "A reconciled income statement flow with dollar and common-size labels."),
            ("How much cash is available after investment?", "A transparent bridge among textbook FCF, cash FCF, and P&G adjusted FCF."),
            ("What changed over time?", "FY2021-FY2025 net sales and GAAP diluted EPS with four-interval CAGRs."),
            ("What drives shareholder return?", "Profit margin x asset turnover x equity multiplier = direct ROE."),
            ("Can we trust the result?", "Source tracing, formula checks, code checks, visual checks, and a documented human correction."),
        ],
        [3000, 6360],
        first_col_bold=True,
    )

    add_section_start(doc, "2. Team Structure and Individual Accountability")
    add_table(
        doc,
        ["Role", "Primary responsibility", "Required cross-check"],
        [
            ("Data Controller", "Confirms source pages, fiscal years, units, signs, and Metric_ID labels.", "Trace one number in every visual to 01_Input and the 10-K."),
            ("Income Statement Lead", "Builds and explains the Sankey and common-size margins.", "Recalculate one FCF subtotal and one DuPont factor."),
            ("Cash Flow Lead", "Builds NOPAT, NOWC, textbook FCF, cash FCF, and the reconciliation.", "Audit a Sankey node and the CAGR exponent."),
            ("Trend and DuPont Lead", "Builds the five-year trend, CAGRs, and ROE decomposition.", "Audit the CapEx sign and adjusted FCF label."),
            ("Python/AI Lead", "Maintains the executable program, prompt history, and REM comments.", "Explain the finance meaning of all four visuals."),
            ("Audit/Presentation Lead", "Coordinates verification, corrections, executive story, and presentation.", "Run the code and trace at least one formula end to end."),
        ],
        [1900, 4200, 3260],
        font_size=8.8,
    )
    add_bullets(doc, [
        "If your team has fewer than six students, combine adjacent roles; do not remove responsibilities.",
        "Every member must be able to locate the workbook input, explain a sign convention, trace a chart label to a source value, and explain why the script stops when a check fails.",
        "The professor may direct an oral question to any team member, regardless of assigned role.",
    ])

    add_section_start(doc, "3. Learning Objectives Tied to Chapters 1-8")
    add_table(
        doc,
        ["Chapter", "Course connection", "Evidence of learning"],
        [
            ("1 - Financial management and value", "Value depends on future cash flows, risk, and decisions that affect owners and creditors.", "Explain why net income alone is not a complete value-creation measure."),
            ("2 - Financial markets", "Public financial reports reduce information gaps between firms and capital providers.", "Use the audited 10-K rather than invented or replacement data."),
            ("3 - Financial statements and FCF", "Income statement, balance sheet, cash flow, taxes, NOPAT, NOWC, and FCF are connected.", "Reconcile the Sankey and the FCF waterfall."),
            ("4 - Financial statement analysis", "Common-size, trend, profitability, efficiency, leverage, and DuPont analysis turn statements into insight.", "Interpret margins, CAGRs, and all three DuPont factors."),
            ("5 - Time value of money", "Growth is measured across time and cash flows matter because timing affects value.", "Calculate CAGR over four intervals and connect FCF to valuation."),
            ("6 - Interest rates", "Financing costs and macro rates affect earnings, investment, and leverage decisions.", "Classify interest after EBIT and discuss financing sensitivity."),
            ("7 - Bonds and creditor risk", "Debt holders care about coverage, cash generation, leverage, and default risk.", "Explain why a high equity multiplier can matter to creditors."),
            ("8 - Risk and return", "Higher expected return generally requires bearing relevant risk; leverage magnifies outcomes.", "Identify operating, financial, and model risks not captured by the dashboard."),
        ],
        [1350, 3900, 4110],
        font_size=8.5,
    )

    add_section_start(doc, "4. Data Rules and Reproducibility")
    add_bullets(doc, [
        "Use PG_FinancialStory_Student.xlsx as the program input. Do not rename 01_Input, its headers, or its Metric_ID values.",
        "Use fiscal years ending June 30. Do not substitute calendar-year, quarterly, or annualized quarterly values.",
        "Use USD millions except per-share data. Label any display conversion to billions.",
        "Preserve raw statement signs. Convert an expense or CapEx to a positive magnitude only when the formula requires it, and document the conversion.",
        "Use consolidated net earnings consistently with total shareholders' equity in the base DuPont model.",
        "Use average FY2024/FY2025 assets and total shareholders' equity for turnover, leverage, and ROE.",
        "Label P&G adjusted FCF as a company-defined non-GAAP measure.",
        "Use relative paths or command-line arguments. The program must run on another computer from the submitted folder.",
    ])
    add_callout(doc, "Source hierarchy", "Audited 10-K and supplied source workbook -> standardized 01_Input values -> team calculations -> Python outputs -> verification log.")

    add_section_start(doc, "5. Required Workflow")
    add_table(
        doc,
        ["Stage", "Team action", "Evidence"],
        [
            ("1. Read", "Locate the statement line items and understand the definitions.", "Source page/tab references."),
            ("2. Calculate", "Complete formulas manually or in the working sheets before charting.", "Formula cells and calculation notes."),
            ("3. Prompt", "Give AI explicit inputs, outputs, constraints, and failure behavior.", "Exact prompt history."),
            ("4. Run", "Execute the program from the submitted folder and save all outputs.", "Executable .py file and output files."),
            ("5. Verify", "Test source values, formulas, code logic, and visual encodings.", "07_AI_Verification and sanity-check CSV."),
            ("6. Revise", "Make at least one meaningful human improvement.", "Before/after description and changed code line/function."),
            ("7. Interpret", "Combine the four lenses into one cautious conclusion.", "Two-page interpretation and five-minute presentation."),
        ],
        [1250, 5000, 3110],
        font_size=9,
    )

    add_section_start(doc, "6. Visualization 1 - Income Statement Sankey")
    doc.add_paragraph(
        "Create a width-proportional flow diagram showing where P&G's FY2025 net sales went. The diagram must preserve the order of the income statement: operating expenses are subtracted before EBIT; interest and other non-operating items are handled after EBIT."
    )
    add_table(
        doc,
        ["Required identity", "FY2025 benchmark"],
        [
            ("Net Sales = COGS + Gross Profit", "84,284 = 41,164 + 43,120"),
            ("Gross Profit = SG&A + Impairment + EBIT", "43,120 = 22,669 + 0 + 20,451"),
            ("EBIT - Net Non-Operating Expense = EBT", "20,451 - 284 = 20,167"),
            ("EBT - Income Taxes = Net Earnings", "20,167 - 4,102 = 16,065"),
        ],
        [5600, 3760],
        first_col_bold=True,
    )
    doc.add_heading("Required labels", level=2)
    add_bullets(doc, [
        "Dollar amount in USD millions or clearly labeled billions.",
        "Percentage of net sales for each major node.",
        "Audited fiscal year and source statement.",
        "Clear operating/non-operating classification.",
        "Flow widths proportional to the same values printed in the labels.",
    ])
    doc.add_heading("What the chart should mean", level=2)
    doc.add_paragraph(
        "For each $1.00 of FY2025 net sales, about $0.49 went to product costs, $0.27 went to SG&A, $0.24 remained as EBIT, and $0.19 remained as consolidated net earnings. Your explanation must distinguish a useful insight from a restatement of labels."
    )
    doc.add_heading("Audit questions", level=2)
    add_bullets(doc, [
        "Does each source node equal the sum of outgoing flows?",
        "Does each target node equal the sum of incoming flows?",
        "Did the program place interest expense after EBIT?",
        "Did the program explain the difference between consolidated net earnings and net earnings attributable to P&G?",
        "Could the colors or labels imply a classification that the accounting does not support?",
    ])

    add_section_start(doc, "7. Visualization 2 - Free Cash Flow Waterfall")
    doc.add_paragraph(
        "Build a waterfall that starts with EBIT and ends with P&G adjusted FCF. Show textbook FCF, cash FCF, and adjusted FCF as separate totals. Do not hide the reconciliation residual."
    )
    add_table(
        doc,
        ["Calculation", "Formula", "FY2025 result"],
        [
            ("Effective tax rate", "Income Taxes / EBT", "20.34%"),
            ("NOPAT", "EBIT x (1 - tax rate)", "$16,291.23M"),
            ("Operating NOWC", "AR + Inventory + Prepaids - AP - Accrued Liabilities", "FY25: -$10,709M; FY24: -$11,208M"),
            ("Change in NOWC", "FY2025 NOWC - FY2024 NOWC", "$499M use of cash"),
            ("Textbook FCF", "NOPAT + D&A - CapEx - change in NOWC", "$14,866.23M"),
            ("Cash FCF", "Cash Flow from Operations - CapEx", "$14,044M"),
            ("Other operating bridge", "Cash FCF - Textbook FCF", "-$822.23M"),
            ("P&G adjusted FCF", "Cash FCF + 2017 Tax Act payment add-back", "$14,606M (non-GAAP)"),
        ],
        [1900, 4600, 2860],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Interpretation",
        "Textbook FCF and cash FCF answer related but different questions. The residual reflects operating accruals, noncash items, and timing effects not fully represented by the simplified NOWC definition. P&G adjusted FCF then applies a company-defined add-back.",
    )
    doc.add_heading("Common AI failure", level=2)
    add_code(doc, "raw_capex = -3773\nwrong_fcf = nopat + da - raw_capex - delta_nowc  # subtracting a negative adds CapEx\ncapex = abs(raw_capex)\nright_fcf = nopat + da - capex - delta_nowc")

    add_section_start(doc, "8. Visualization 3 - Five-Year Revenue and EPS Trend")
    doc.add_paragraph(
        "Create one visualization for FY2021-FY2025 using net sales and GAAP diluted EPS. The units differ, so the chart must make both scales explicit and avoid implying that the two series have identical economic meaning."
    )
    add_table(
        doc,
        ["Fiscal year", "Net sales ($M)", "GAAP diluted EPS"],
        [
            ("FY2021", "76,118", "$5.50"),
            ("FY2022", "80,187", "$5.81"),
            ("FY2023", "82,006", "$5.90"),
            ("FY2024", "84,039", "$6.02"),
            ("FY2025", "84,284", "$6.51"),
        ],
        [2200, 3500, 3660],
    )
    add_code(doc, "Sales CAGR = (84,284 / 76,118) ** (1 / 4) - 1 = 2.58%\nEPS CAGR   = (6.51 / 5.50) ** (1 / 4) - 1 = 4.30%")
    add_bullets(doc, [
        "Use four compounding intervals, not five observations, in the exponent.",
        "State that EPS grew faster than net sales over the period.",
        "Evaluate possible drivers: operating margins, impairment charges, taxes, non-operating items, diluted shares, and mix.",
        "Do not mix GAAP diluted EPS with core or adjusted EPS.",
        "Explain how a dual-axis chart could be numerically correct yet visually misleading.",
    ])

    add_section_start(doc, "9. Visualization 4 - DuPont ROE Decomposition")
    doc.add_paragraph(
        "Use the simplified consolidated DuPont identity with average FY2024/FY2025 assets and total shareholders' equity. Keep numerator and denominator definitions internally consistent."
    )
    add_code(doc, "ROE = Profit Margin x Total Asset Turnover x Equity Multiplier")
    add_table(
        doc,
        ["Factor", "Formula", "P&G result"],
        [
            ("Profit margin", "16,065 / 84,284", "19.06%"),
            ("Total asset turnover", "84,284 / average assets of 123,800.5", "0.681x"),
            ("Equity multiplier", "average assets of 123,800.5 / average equity of 51,421.5", "2.408x"),
            ("DuPont ROE", "0.1906 x 0.6808 x 2.4076", "31.24%"),
            ("Direct ROE check", "16,065 / 51,421.5", "31.24%"),
        ],
        [2000, 5000, 2360],
        font_size=8.9,
    )
    add_bullets(doc, [
        "Explain why average balance-sheet values match a full-year income measure better than ending values.",
        "Describe P&G's strong profit margin, relatively modest asset turnover, and substantial equity multiplier.",
        "Explain why leverage can raise ROE when performance is favorable and magnify risk when it is not.",
        "Confirm that the three-factor product equals direct ROE within tolerance.",
    ])

    doc.add_page_break()
    add_section_start(doc, "10. AI Requirements")
    add_bullets(doc, [
        "AI use is required. Submit the exact prompt or prompt sequence used for the final program.",
        "AI may draft code and explanations; it may not replace the supplied audited data with searched, estimated, or invented values.",
        "Your final program must be executable, commented, reproducible, and organized into meaningful functions.",
        "Use Metric_ID values instead of fixed row numbers.",
        "The program must validate required columns, metrics, years, and accounting identities.",
        "The program must stop with a clear error if a required identity fails.",
        "Every team member must be able to explain the code, formulas, labels, and limitations.",
        "Document one meaningful human revision after the first AI draft.",
    ])
    add_callout(doc, "What REM means here", "REM is a visible, concise record of reasoning summary, equations, method, and limitations. It is not a request for private chain-of-thought.")

    add_section_start(doc, "11. Required Model AI Prompt")
    add_code(doc, """You are assisting a FINC 310 student team. Use the attached
PG_FinancialStory_Student.xlsx workbook. Do not search for or invent replacement
financial data.

Write one executable Python program named pg_financial_story.py that:
1. reads the 01_Input worksheet using the Metric_ID column;
2. validates that required metrics and fiscal years exist;
3. creates an FY2025 income statement Sankey;
4. creates a free cash flow waterfall that separately labels textbook FCF,
   cash FCF, and P&G adjusted FCF;
5. creates an FY2021-FY2025 net sales and GAAP diluted EPS trend;
6. creates a DuPont ROE decomposition using average FY2024/FY2025 assets and
   total shareholders' equity;
7. writes a CSV of accounting and finance sanity checks;
8. stops with a clear error message if a required identity fails;
9. saves four individual PNG files and one combined dashboard PNG; and
10. uses relative paths or command-line arguments.

Begin the file with concise REM comments containing Reasoning Summary,
Equations, Method, and Limitations. Comment every major function. Use USD
millions except EPS. Explain all sign conventions. Do not provide hidden
chain-of-thought; provide concise auditable rationale. After the code, list the
expected output files and the commands needed to install packages and run it.""")
    doc.add_heading("Required follow-up prompts", level=2)
    add_numbered(doc, [
        "Audit your own code for hard-coded financial values that should be read from the workbook. Return a table of every hard-coded number and whether it is justified.",
        "List every accounting identity used in the program and show where the program tests it.",
        "Explain how the code handles negative interest expense and negative capital expenditures.",
        "Identify three ways the charts could be numerically correct but visually misleading.",
        "Rewrite any function that depends on fixed Excel row numbers so that it uses Metric_ID values.",
        "Generate five oral questions the professor could ask, then recommend one meaningful code revision.",
    ])

    add_section_start(doc, "12. Required REM Statements")
    add_code(doc, """# REM - REASONING SUMMARY
# Explain the program's purpose and how the four visuals fit together.

# REM - EQUATIONS
# List Gross Profit, NOPAT, Operating NOWC, textbook FCF, cash FCF,
# CAGR, Profit Margin, Asset Turnover, Equity Multiplier, and ROE.

# REM - METHOD
# Explain workbook input, validation, calculations, chart creation,
# output files, and failure handling.

# REM - LIMITATIONS
# Explain simplified NOWC, non-GAAP adjusted FCF, fiscal-year scope,
# source dependence, rounding, and risks of misleading chart design.""")
    add_table(
        doc,
        ["REM section", "Minimum acceptable content"],
        [
            ("Reasoning Summary", "What the program is designed to accomplish and how the outputs answer different financial questions."),
            ("Equations", "Every financial formula, sign convention, denominator choice, and CAGR interval."),
            ("Method", "Input, validation, model construction, plotting, export, and failure behavior."),
            ("Limitations", "Simplifications, non-GAAP definitions, source scope, rounding, visual design risk, and missing forward-looking risk."),
        ],
        [2400, 6960],
        first_col_bold=True,
    )

    doc.add_page_break()
    add_section_start(doc, "13. Annotated Python Starter Architecture")
    doc.add_paragraph(
        "The scaffold below is not a complete solution. It shows the architecture and control points your final program must contain. Annotation labels A1-A12 explain what each block is responsible for."
    )
    add_code(doc, """from pathlib import Path
import math
import pandas as pd

# A1 - Read the standardized sheet by header names, not fixed cell addresses.
def load_metrics(workbook_path: Path) -> dict[str, dict[int, float]]:
    df = pd.read_excel(workbook_path, sheet_name="01_Input", header=4)
    required = {"Metric_ID", "FY2025", "FY2024", "FY2023", "FY2022", "FY2021"}
    missing = required.difference(df.columns)
    if missing:
        raise ValueError(f"Missing required columns: {sorted(missing)}")

    # A2 - Metric_ID makes the program resistant to inserted or moved rows.
    df = df[df["Metric_ID"].notna()].copy()
    metrics = {}
    for _, row in df.iterrows():
        metric_id = str(row["Metric_ID"]).strip()
        metrics[metric_id] = {}
        for year in (2025, 2024, 2023, 2022, 2021):
            value = row.get(f"FY{year}")
            if pd.notna(value):
                metrics[metric_id][year] = float(value)
    return metrics

# A3 - One accessor produces a clear failure instead of a silent blank or zero.
def value(metrics, metric_id: str, year: int) -> float:
    try:
        return metrics[metric_id][year]
    except KeyError as exc:
        raise KeyError(f"Missing {metric_id} for FY{year}") from exc

def build_model(metrics) -> dict[str, float]:
    model = {}
    model["sales"] = value(metrics, "NET_SALES", 2025)
    model["cogs"] = value(metrics, "COGS", 2025)
    model["gross_profit"] = model["sales"] - model["cogs"]  # A4 - identity

    # A5 - Raw CapEx is negative on the cash-flow statement; the formula uses a
    # positive magnitude that is subtracted once.
    model["capex"] = abs(value(metrics, "CAPEX", 2025))

    # A6 - Use average balance-sheet values with full-year sales and earnings.
    model["avg_assets"] = (
        value(metrics, "TOTAL_ASSETS", 2025)
        + value(metrics, "TOTAL_ASSETS", 2024)
    ) / 2

    # A7 - CAGR has four intervals between five annual observations.
    sales_2021 = value(metrics, "NET_SALES", 2021)
    model["sales_cagr"] = (model["sales"] / sales_2021) ** (1 / 4) - 1
    return model

# A8 - Checks compare two independently constructed sides of each identity.
def close_enough(actual: float, expected: float, tolerance: float) -> bool:
    return math.isclose(actual, expected, abs_tol=tolerance)

def require_check(name: str, actual: float, expected: float, tolerance: float) -> None:
    if not close_enough(actual, expected, tolerance):
        raise RuntimeError(f"{name} failed: {actual} versus {expected}")  # A9

def main() -> int:
    workbook = Path("PG_FinancialStory_Student.xlsx")  # A10 - relative path
    metrics = load_metrics(workbook)
    model = build_model(metrics)
    require_check(
        "Revenue identity",
        model["cogs"] + model["gross_profit"],
        model["sales"],
        1.0,
    )
    # A11 - Call one function per chart only after all required checks pass.
    # save_sankey(model, Path("PG_01_Income_Statement_Sankey.png"))
    # save_waterfall(...); save_trend(...); save_dupont(...)
    return 0

if __name__ == "__main__":
    raise SystemExit(main())  # A12 - returns a visible process status.""")
    add_table(
        doc,
        ["Annotation", "What you must be able to explain"],
        [
            ("A1-A3", "How the program reads the workbook and fails when inputs are missing."),
            ("A4-A7", "Where finance definitions, sign rules, averages, and time intervals enter the model."),
            ("A8-A9", "Why tolerance-based checks are stronger than accepting printed output."),
            ("A10-A12", "Why relative paths, ordered execution, and process status support reproducibility."),
        ],
        [1800, 7560],
        first_col_bold=True,
    )

    add_section_start(doc, "14. AI Verification Requirements")
    add_table(
        doc,
        ["Layer", "Question", "Required evidence"],
        [
            ("Source", "Did the program read the correct audited value, period, unit, and sign?", "Workbook cell/Metric_ID and 10-K page/tab."),
            ("Formula", "Does the equation match the finance definition and denominator choice?", "Manual calculation and formula worksheet."),
            ("Program", "Does code implement the formula without hard-coded or sign errors?", "Annotated function and test result."),
            ("Visual", "Do widths, directions, axes, labels, and colors tell the same truth as the data?", "Chart audit and corrected output."),
            ("Interpretation", "Does the narrative distinguish evidence, inference, and limitation?", "Cautious written conclusion."),
        ],
        [1350, 4100, 3910],
        font_size=8.9,
    )
    doc.add_heading("Six required automatic checks", level=2)
    add_bullets(doc, [
        "Revenue identity: COGS + Gross Profit = Net Sales.",
        "Gross profit identity: SG&A + Impairment + EBIT = Gross Profit.",
        "Pretax identity: EBIT - Net Non-Operating Expense = EBT.",
        "Net earnings identity: EBT - Income Taxes = Net Earnings.",
        "Adjusted FCF identity: Cash FCF + Tax Act payment add-back = P&G Adjusted FCF.",
        "DuPont identity: product of the three factors = direct ROE.",
    ])
    doc.add_heading("Required human audit narrative", level=2)
    add_numbered(doc, [
        "Identify one AI-generated assumption that required human judgment.",
        "Identify one place where the code could silently produce a misleading chart.",
        "Describe one numerical, code, or labeling mistake your team corrected.",
        "Explain why P&G adjusted FCF differs from textbook FCF.",
        "State whether the final dashboard is reasonable and cite at least three checks.",
    ])

    doc.add_page_break()
    add_section_start(doc, "15. Required Deliverables")
    add_table(
        doc,
        ["Deliverable", "File or format", "Acceptance standard"],
        [
            ("Completed workbook", "PG_FinancialStory_Student.xlsx", "Yellow response/calculation cells complete; formulas and sources retained."),
            ("Executable program", "pg_financial_story.py", "Runs from the submitted folder with relative/command-line paths."),
            ("Four visuals", "Four PNG files", "Required labels, proportional encodings, and source notes."),
            ("Combined dashboard", "One PNG", "All four visuals are legible on one page."),
            ("Sanity-check log", "CSV plus 07_AI_Verification", "All required checks pass with formulas and evidence."),
            ("AI documentation", "Prompt history and REM block", "Exact prompts, equations, method, limitations, and follow-ups."),
            ("Annotated code", "Complete printed code or commented PDF/Word", "Inputs, calculations, outputs, and failure risk identified for each major function."),
            ("Financial interpretation", "Two pages maximum", "One integrated story with evidence, inference, and limitation."),
            ("Presentation", "Five minutes", "Every member participates and answers questions."),
        ],
        [1900, 2800, 4660],
        font_size=8.6,
    )
    add_callout(doc, "File test", "Before submitting, copy the package to a different folder and run the exact command another person would use.", fill=GREEN, accent="25633D")

    add_section_start(doc, "16. Presentation Expectations")
    add_table(
        doc,
        ["Time", "Content", "Required message"],
        [
            ("0:00-0:30", "Thesis", "State the most important conclusion from the dashboard."),
            ("0:30-1:30", "Income statement", "Explain the sales-dollar flow and one audited identity."),
            ("1:30-2:30", "Free cash flow", "Distinguish textbook, cash, and adjusted FCF; explain the residual."),
            ("2:30-3:30", "Trend and DuPont", "Explain the growth gap and the profitability/efficiency/leverage mix."),
            ("3:30-4:30", "AI audit", "Show one correction, one automatic check, and one visual risk."),
            ("4:30-5:00", "Caveat and close", "Name one limitation and return to the thesis."),
        ],
        [1500, 2600, 5260],
        font_size=9,
    )
    add_bullets(doc, [
        "Use no more than five content slides plus an optional title slide.",
        "Every team member must speak and be ready for a random question.",
        "Do not read code line by line. Explain the financial job performed by the code.",
        "Use large labels; the four figures must be readable from the back of the room.",
        "State when a conclusion is an inference rather than a fact directly shown by the data.",
        "Finish within five minutes; a concise, verified story is stronger than a rushed inventory of details.",
    ])

    add_section_start(doc, "17. Quality and Grading Priorities")
    add_table(
        doc,
        ["Priority", "What earns credit", "What loses credit"],
        [
            ("1. Accuracy and provenance", "Correct fiscal year, units, signs, audited sources, and reconciliations.", "Invented data, wrong period, or material unreconciled values."),
            ("2. Financial understanding", "Definitions and interpretations are internally consistent.", "Treating textbook FCF and adjusted FCF as interchangeable."),
            ("3. Verification", "Automatic checks plus documented human review and correction.", "Saying 'AI was correct' without evidence."),
            ("4. Reproducibility", "Code runs from submitted files on another computer.", "Hard-coded personal paths or unexplained manual steps."),
            ("5. Communication", "One integrated, cautious, readable story.", "Attractive but misleading axes, widths, labels, or claims."),
        ],
        [2100, 3630, 3630],
        font_size=8.8,
    )

    add_section_start(doc, "Appendix A - Formula and Benchmark Sheet")
    add_table(
        doc,
        ["Item", "Formula or benchmark"],
        [
            ("Gross Profit", "Net Sales - Cost of Products Sold = $43,120M"),
            ("Net Non-Operating Expense", "$907M interest expense - $469M interest income - $154M other income = $284M"),
            ("Effective tax rate", "$4,102M / $20,167M = 20.34%"),
            ("NOPAT", "$20,451M x (1 - 20.34%) = $16,291.23M"),
            ("Textbook FCF", "$16,291.23M + $2,847M - $3,773M - $499M = $14,866.23M"),
            ("Cash FCF", "$17,817M - $3,773M = $14,044M"),
            ("P&G adjusted FCF", "$14,044M + $562M = $14,606M (non-GAAP)"),
            ("Sales CAGR", "2.58% over four intervals"),
            ("EPS CAGR", "4.30% over four intervals"),
            ("DuPont ROE", "19.06% x 0.6808x x 2.4076x = 31.24%"),
        ],
        [2800, 6560],
        first_col_bold=True,
        font_size=9,
    )

    add_section_start(doc, "Appendix B - Oral Readiness Questions")
    add_bullets(doc, [
        "Why is interest expense excluded from EBIT?",
        "Why is depreciation and amortization added back in textbook FCF?",
        "Why does an increase in operating NOWC reduce FCF?",
        "Why can textbook FCF differ from CFO minus CapEx without either measure being automatically wrong?",
        "Why is P&G adjusted FCF non-GAAP?",
        "Why does CAGR use a one-fourth exponent across five annual observations?",
        "Why use average assets and equity in DuPont analysis?",
        "How can a high equity multiplier raise both ROE and risk?",
        "Where does the program convert a raw negative CapEx to a positive magnitude?",
        "Which statement in the program prevents a bad chart from being exported?",
        "What was your team's meaningful human correction?",
        "What important risk or forward-looking factor is missing from the dashboard?",
    ])

    path = OUT_DIR / "PG_Student_Group_Assignment_Complete.docx"
    doc.save(path)
    return path


def build_instructor_guide() -> Path:
    doc = Document()
    style_document(doc, "FINC 310 | P&G Financial Story | Instructor Teaching Guide")
    add_cover(
        doc,
        "FINC 310 Instructor Guide",
        "Teaching the P&G Financial Story",
        "A flexible 2-4 hour classroom plan for financial concepts, AI-assisted Python, verification, and student presentations.",
        [
            ("Audience", "FINC 310 students after or during Chapters 1-8"),
            ("Core version", "120 minutes: concepts, four visuals, AI prompt, and verification demo"),
            ("Lab version", "180 minutes: core plus guided Python/AI work and audit challenge"),
            ("Full version", "240 minutes: lab plus team presentations and peer critique"),
            ("Files", "Student workbook, 10-K, student assignment, classroom deck, rubric, and solution package"),
        ],
    )

    add_section_start(
        doc,
        "Instructional Premise",
        "Treat AI as a junior analyst whose work must be independently verified, not as an answer key.",
    )
    doc.add_paragraph(
        "The class moves from source evidence to finance meaning, then to code and verification. Students first learn what the four visuals mean and which accounting identities constrain them. Only then do they use AI to draft code. This sequence keeps the financial model in charge of the technology."
    )
    add_callout(
        doc,
        "Communication job",
        "By the end, students should be able to build and defend a reproducible P&G financial story because they can connect audited statements to formulas, visual encodings, code, and verification evidence.",
    )
    add_table(
        doc,
        ["Design principle", "Classroom implication"],
        [
            ("Evidence before output", "Open the 10-K and workbook before showing code."),
            ("Meaning before syntax", "Explain the finance job of each function before reviewing Python details."),
            ("Checks before charts", "A failed identity stops export; students do not 'fix' a chart by hiding a residual."),
            ("Roles without silos", "Teams distribute labor, but oral questions enforce whole-team understanding."),
            ("AI with accountability", "Prompt history, REM comments, corrections, and verification are graded evidence."),
        ],
        [3000, 6360],
        first_col_bold=True,
    )

    add_section_start(doc, "1. Learning Outcomes and Course Alignment")
    add_table(
        doc,
        ["Outcome", "Student evidence", "Chapter connection"],
        [
            ("Explain value creation beyond accounting profit.", "Connect net income, FCF, risk, and shareholder value.", "Ch. 1"),
            ("Use public-company information responsibly.", "Trace metrics to the audited 10-K and standardized workbook.", "Ch. 2"),
            ("Link statements, taxes, NOPAT, NOWC, and FCF.", "Reconcile the Sankey and FCF waterfall.", "Ch. 3"),
            ("Analyze profitability, efficiency, leverage, and growth.", "Interpret common-size margins, trend CAGRs, and DuPont.", "Ch. 4"),
            ("Reason across time.", "Use four CAGR intervals and connect future FCF to present value.", "Ch. 5"),
            ("Recognize financing sensitivity.", "Place interest after EBIT and discuss rate/leverage effects.", "Ch. 6"),
            ("Adopt a creditor perspective.", "Explain cash generation, coverage, leverage, and default implications.", "Ch. 7"),
            ("Distinguish return from risk.", "Explain why leverage magnifies ROE and financial risk.", "Ch. 8"),
            ("Control AI-assisted analysis.", "Audit sources, formulas, code, visuals, and interpretations.", "Across course"),
        ],
        [3000, 4350, 2010],
        font_size=8.6,
    )

    add_section_start(doc, "2. Instructor Preparation")
    doc.add_heading("Before class", level=2)
    add_bullets(doc, [
        "Post PG_Student_Group_Assignment_Complete.docx, PG_FinancialStory_Student.xlsx, the FY2025 annual report, requirements.txt, and the grading rubric.",
        "Keep the instructor workbook, solution manual, solution Python, and finished visuals unpublished unless you plan to reveal them during class.",
        "Run the solution program against the student workbook and confirm the six sanity checks pass.",
        "Open the classroom deck in Presenter View so timing and facilitation notes remain visible to you.",
        "Decide whether students will install Python locally, use a managed lab, or watch a live demo and complete coding outside class.",
        "Form teams of approximately six and assign roles before the lab segment.",
        "Prepare a backup folder containing the verified output images in case local Python installation fails.",
    ])
    doc.add_heading("Room and technology", level=2)
    add_table(
        doc,
        ["Need", "Preferred setup", "Fallback"],
        [
            ("Display", "Projector with 16:9 PowerPoint and Presenter View.", "PDF/printed slide thumbnails."),
            ("Student devices", "One laptop per team with Python and Excel.", "Instructor demo plus paper audit activities."),
            ("AI access", "Approved AI tool capable of handling prompt text and code.", "Provide a printed AI draft for critique."),
            ("Files", "One shared folder per team.", "USB or LMS download package."),
            ("Submission", "One group upload plus individual oral checks.", "Exit ticket plus later group upload."),
        ],
        [1500, 4200, 3660],
        font_size=9,
    )
    doc.add_heading("Student prework", level=2)
    add_bullets(doc, [
        "Read the 00_ReadMe worksheet and inspect the 01_Input Metric_ID structure.",
        "Bring the student assignment and workbook to class.",
        "Install packages with: python -m pip install -r requirements.txt.",
        "Write a one-sentence prediction: Will P&G's EPS growth be faster or slower than revenue growth, and why?",
    ])

    add_section_start(doc, "3. Master 240-Minute Run of Show")
    add_table(
        doc,
        ["Minutes", "Slides", "Learning move", "Student evidence"],
        [
            ("0-10", "1-4", "Open with the earnings-versus-sales question; reveal the four-lens dashboard.", "Initial prediction and one question."),
            ("10-25", "5-8", "Orient teams to the assignment, roles, source hierarchy, and story map.", "Role assignments and source trace."),
            ("25-50", "9-12", "Teach income-statement flow; run the Sankey classification/reconciliation activity.", "Balanced node and classification correction."),
            ("50-80", "13-18", "Contrast profit and cash; build the FCF bridge; diagnose the CapEx sign error.", "Correct sign logic and residual explanation."),
            ("80-90", "Break", "Pause and reset.", "None."),
            ("90-110", "19-22", "Teach multi-year trend, four-interval CAGR, and dual-axis risk.", "Growth-gap hypothesis."),
            ("110-135", "23-26", "Teach DuPont, average balances, and leverage risk.", "Factor interpretation and direct-ROE check."),
            ("135-155", "27-30", "Frame AI as junior analyst; unpack prompt structure and REM documentation.", "Team prompt revision."),
            ("155-180", "31-34", "Walk through input, model, sign, and sanity-check code architecture.", "Annotated code trace."),
            ("180-205", "35-37", "Run the four-layer verification lab and document a human correction.", "Verification log and correction statement."),
            ("205-235", "38-41", "Team presentations and peer critique using the accuracy-first rubric.", "Five-minute story and one peer question."),
            ("235-240", "42-43", "Launch next steps and collect exit tickets.", "One formula, one check, one risk."),
        ],
        [1100, 900, 4800, 2560],
        font_size=8.2,
    )

    add_section_start(doc, "4. Compression Paths for 2 or 3 Hours")
    doc.add_heading("120-minute core class", level=2)
    add_table(
        doc,
        ["Minutes", "Focus", "Slides and cuts"],
        [
            ("0-15", "Case, outcomes, assignment, and source", "1-8; shorten role discussion."),
            ("15-35", "Sankey concept and audit", "9-12; complete one node live."),
            ("35-60", "FCF bridge and sign error", "13-18; emphasize three FCF definitions."),
            ("60-65", "Micro-break", "Pause."),
            ("65-80", "Five-year trend", "19-22; one-minute pair discussion."),
            ("80-95", "DuPont", "23-26; focus on averages and leverage."),
            ("95-115", "AI prompt, REM, code, and checks", "27-36; demo rather than team lab."),
            ("115-120", "Exit ticket", "43; assign 37-42 as outside work."),
        ],
        [1200, 3650, 4510],
        font_size=8.8,
    )
    doc.add_heading("180-minute lab class", level=2)
    add_bullets(doc, [
        "Use the first 135 minutes of the master plan through DuPont.",
        "Use 35 minutes for AI prompt, REM, and annotated code (slides 27-34).",
        "Use the final 10 minutes for verification setup and an exit ticket; presentations occur next class or by recorded submission.",
        "If students already know Chapters 3-4 well, reduce the concepts segment to 105 minutes and reserve 55 minutes for hands-on coding and verification.",
    ])

    add_section_start(doc, "5. Facilitation Guide by Module")
    doc.add_heading("Opening and source orientation - Slides 1-8", level=2)
    add_bullets(doc, [
        "Ask: 'Can EPS grow faster than sales without strong revenue growth?' Record two competing explanations before revealing numbers.",
        "Show the dashboard only long enough to create questions. Do not explain each visual yet.",
        "Open the annual report's Consolidated Statements of Earnings and have students locate net sales, operating income, net earnings, and diluted EPS.",
        "Emphasize that public markets depend on credible information; the 10-K is evidence, not decoration.",
        "Assign roles and state that roles divide labor but not accountability.",
    ])
    add_callout(doc, "Check for understanding", "Ask one student to trace $84,284M from the 10-K to Metric_ID NET_SALES in 01_Input and then to a chart label.")

    doc.add_heading("Income statement Sankey - Slides 9-12", level=2)
    add_bullets(doc, [
        "Build the statement in four identities before showing the finished Sankey.",
        "Ask students to calculate each major item as a percentage of net sales. Translate the result into cents per sales dollar.",
        "Point out that a Sankey is constrained by conservation: every source node must balance.",
        "Use the interest-expense classification error to test whether students understand EBIT rather than merely reading labels.",
        "Ask why consolidated net earnings and earnings attributable to P&G differ by $91M and why numerator/denominator consistency matters.",
    ])
    add_callout(doc, "Expected insight", "Product cost absorbs 48.8% of sales, SG&A 26.9%, EBIT is 24.3%, and consolidated net earnings are 19.1% of sales.")

    doc.add_heading("Free cash flow waterfall - Slides 13-18", level=2)
    add_bullets(doc, [
        "Begin with the conceptual contrast: income statement profit includes accrual accounting; FCF asks what cash remains after operating investment.",
        "Build NOPAT, operating NOWC, change in NOWC, and textbook FCF sequentially.",
        "Then introduce cash FCF from the cash flow statement and treat the $822.23M residual as a bridge to explain.",
        "Finally reveal the $562M Tax Act payment add-back and require the non-GAAP label.",
        "Demonstrate the double-negative CapEx bug and ask students to predict the direction of error before calculating it.",
    ])
    add_callout(doc, "Expected insight", "Textbook FCF is $14.866B, cash FCF is $14.044B, and P&G adjusted FCF is $14.606B. Similar totals do not make the definitions interchangeable.")

    doc.add_heading("Five-year trend - Slides 19-22", level=2)
    add_bullets(doc, [
        "Have students count the spaces between FY2021 and FY2025: there are four compounding intervals.",
        "Discuss why a dual-axis chart can exaggerate or minimize co-movement through arbitrary axis bounds.",
        "Ask teams to generate at least three possible explanations for EPS growing faster than sales, then classify each as supported, plausible, or untested.",
        "Keep GAAP diluted EPS distinct from core or adjusted EPS.",
    ])
    add_callout(doc, "Expected insight", "Sales CAGR is 2.58%; diluted EPS CAGR is 4.30%. The chart supports a growth gap, not a single proven cause.")

    doc.add_heading("DuPont ROE - Slides 23-26", level=2)
    add_bullets(doc, [
        "Write ROE as three factors before inserting any numbers.",
        "Use average assets and equity because sales and earnings measure a full period while the balance sheet is a point-in-time snapshot.",
        "Interpret the 19.06% margin, 0.681x turnover, and 2.408x equity multiplier separately.",
        "Avoid declaring one factor 'the cause' of ROE without qualification; say margin supports ROE and leverage magnifies it while turnover is relatively modest.",
        "Connect leverage to Chapter 7 creditor concerns and Chapter 8 risk-return tradeoffs.",
    ])
    add_callout(doc, "Expected insight", "0.1906 x 0.6808 x 2.4076 = 0.3124, matching direct ROE of 16,065 / 51,421.5 = 31.24%.")

    doc.add_heading("AI, code, and verification - Slides 27-37", level=2)
    add_bullets(doc, [
        "Ask students what a junior analyst should be told before beginning: approved data, definitions, outputs, constraints, and failure behavior.",
        "Reveal the model prompt and identify each control. Have teams add one missing or stronger control.",
        "Clarify that REM comments are auditable documentation, not hidden chain-of-thought.",
        "Walk the architecture in financial order: load -> model -> check -> visualize -> export.",
        "At every code slide, ask two questions: What finance job does this block perform? How could it fail silently?",
        "Use the four-layer verification model: source, formula, program, visual, plus interpretation as the narrative layer.",
        "Require a meaningful correction such as removing a fixed row, correcting a sign, improving validation, clarifying a label, or preventing a misleading axis.",
    ])

    add_section_start(doc, "6. Activity Instructions and Answer Keys")
    doc.add_heading("Activity 1 - Sankey classification audit", level=2)
    doc.add_paragraph("Prompt: A draft subtracts interest expense with SG&A before calculating EBIT. Is the chart acceptable? Correct the sequence and state the effect on EBIT.")
    add_callout(doc, "Answer", "No. Interest expense is non-operating in this statement and belongs after EBIT. Moving it before EBIT understates operating income and destroys the meaning of the EBIT node.", fill=GREEN, accent="25633D")

    doc.add_heading("Activity 2 - CapEx sign bug", level=2)
    doc.add_paragraph("Prompt: The workbook stores CapEx as -3,773. AI writes FCF = NOPAT + D&A - raw_capex - delta_nowc. Diagnose the error.")
    add_callout(doc, "Answer", "Subtracting raw_capex subtracts a negative and adds $3,773M. Convert to a positive magnitude once with abs(), then subtract it once. Alternatively preserve the negative sign and add raw_capex, but document the convention consistently.", fill=GREEN, accent="25633D")

    doc.add_heading("Activity 3 - Growth-gap claims", level=2)
    doc.add_paragraph("Prompt: EPS CAGR exceeds sales CAGR. Which explanations are facts, plausible hypotheses, or unsupported claims?")
    add_table(
        doc,
        ["Statement", "Classification", "Teaching point"],
        [
            ("EPS CAGR was 4.30% and sales CAGR was 2.58%.", "Fact", "Directly calculated from consistent GAAP values."),
            ("Margins, taxes, impairment charges, or share count may explain the gap.", "Plausible hypothesis", "Requires supporting analysis before claiming causation."),
            ("P&G's brands became stronger every year.", "Unsupported", "The chart does not measure brand strength."),
            ("EPS growth proves value creation.", "Incomplete", "Value also depends on cash flow, timing, risk, and required return."),
        ],
        [3750, 1900, 3710],
        font_size=8.7,
    )

    doc.add_heading("Activity 4 - DuPont strength and risk", level=2)
    doc.add_paragraph("Prompt: Which factor supports P&G's ROE, which factor constrains it, and what risk is visible?")
    add_callout(doc, "Answer", "The 19.06% profit margin strongly supports ROE. The 0.681x asset turnover is comparatively modest. The 2.408x equity multiplier magnifies the product to 31.24% and signals financial leverage, which can magnify downside as well as upside.", fill=GREEN, accent="25633D")

    doc.add_heading("Activity 5 - Visual verification", level=2)
    doc.add_paragraph("Give teams a chart that is numerically correct but uses a truncated axis, mismatched widths, vague units, or colors that imply a false classification. Ask them to write a correction request to the AI.")
    add_callout(doc, "Strong correction prompt", "Keep the verified values unchanged. Start the sales axis at zero, label both y-axes and units, use width-proportional Sankey flows from the displayed dollar amounts, and distinguish GAAP from non-GAAP measures in the title and note.", fill=GREEN, accent="25633D")

    add_section_start(doc, "7. Live AI/Python Demonstration Script")
    add_numbered(doc, [
        "Open the prompt and underline the data lock: use the supplied workbook; do not invent replacements.",
        "Underline the output contract: four PNGs, one dashboard, one CSV, one executable program.",
        "Underline the failure contract: validate inputs and stop when an identity fails.",
        "Ask AI to draft only load_metrics() first. Inspect its dependence on header names, sheet names, row numbers, and missing values.",
        "Ask AI to list all hard-coded numbers and justify each. Remove any hard-coded financial statement value.",
        "Add build_model() one formula family at a time. Test signs after CapEx and interest logic.",
        "Add sanity_checks() before any plotting function. Intentionally change NET_SALES by $10M and demonstrate that export stops.",
        "Restore the verified input, run the program, and compare the CSV to PG_Sanity_Checks.csv.",
        "Inspect one chart for visual truth: axis bounds, width proportionality, units, fiscal year, and GAAP/non-GAAP labels.",
        "Record a human correction and ask a student to explain the revised block without AI assistance.",
    ])
    add_code(doc, "python -m pip install -r requirements.txt\npython pg_financial_story.py --workbook PG_FinancialStory_Student.xlsx --output PG_Output")

    add_section_start(doc, "8. Assessment Plan")
    add_table(
        doc,
        ["When", "Assessment", "Evidence", "Response"],
        [
            ("Opening", "Prediction", "Reason EPS may differ from sales.", "Surface prior knowledge; do not grade accuracy."),
            ("After Sankey", "Node audit", "One balanced identity and classification.", "Correct EBIT/non-operating confusion immediately."),
            ("After FCF", "Sign check", "Correct CapEx logic and residual explanation.", "Rebuild the formula if students rely on memorized signs."),
            ("After trend", "Claim sorting", "Fact vs hypothesis vs unsupported claim.", "Require evidence language."),
            ("After DuPont", "Factor interpretation", "Strength, constraint, and risk.", "Tie leverage to Chapters 7-8."),
            ("AI lab", "Code trace", "Input -> formula -> check -> chart label.", "Ask a different team member to explain each link."),
            ("End", "Exit ticket", "One formula, one check, one risk.", "Use gaps to plan the next class."),
            ("Submission", "100-point rubric", "Files, checks, interpretation, presentation, oral questions.", "Adjust individual score if understanding is not demonstrated."),
        ],
        [1200, 2100, 3100, 2960],
        font_size=8.3,
    )

    add_section_start(doc, "9. Common Misconceptions and Instructor Responses")
    add_table(
        doc,
        ["Misconception", "Diagnostic question", "Instructor response"],
        [
            ("Net income equals cash generated.", "Where do D&A, CapEx, and working capital enter?", "Contrast accrual earnings with operating and investing cash flows."),
            ("All FCF definitions should match.", "What is the starting point and definition of each measure?", "Require labels and reconciliation instead of forced equality."),
            ("Negative NOWC is automatically bad.", "How do supplier terms and customer collection affect NOWC?", "Focus on business model and change, not sign alone."),
            ("Five years means five CAGR periods.", "How many spaces are between FY2021 and FY2025?", "Count four compounding intervals."),
            ("High ROE always means excellent operations.", "What happens when the equity multiplier rises?", "Separate profitability, efficiency, and leverage."),
            ("A PASS cell proves correctness.", "Who verified the expected value and formula?", "A check is only as good as its independent construction."),
            ("AI-generated code is self-explanatory.", "Can you trace one label to the workbook?", "Require code annotation and oral explanation."),
            ("A beautiful chart is a truthful chart.", "Could axis or width choices mislead?", "Audit visual encoding separately from numerical accuracy."),
        ],
        [2900, 2800, 3660],
        font_size=8.2,
    )

    doc.add_page_break()
    add_section_start(doc, "10. Technical Troubleshooting")
    add_table(
        doc,
        ["Symptom", "Likely cause", "Resolution"],
        [
            ("ModuleNotFoundError", "Packages not installed in the active Python environment.", "Run python -m pip install -r requirements.txt using the same python command used to run the script."),
            ("Workbook not found", "Personal absolute path or wrong working folder.", "Use relative paths/arguments and run from the submission folder."),
            ("Missing expected columns", "Wrong header row, renamed sheet, or edited headers.", "Restore 01_Input structure and read with header=4."),
            ("Missing metric/year", "Blank or renamed Metric_ID or wrong year.", "Validate the exact ID and fiscal-year column; do not replace with zero."),
            ("FCF is too high", "Negative CapEx subtracted twice.", "Use one documented sign convention."),
            ("DuPont does not equal direct ROE", "Mixed net income/equity definitions or ending balances.", "Use consolidated earnings and average total equity consistently."),
            ("Chart saves but labels are wrong", "Hard-coded label text or units.", "Build labels from the model dictionary and test them."),
            ("Excel shows stale formulas", "Workbook formula cache not recalculated.", "Open and calculate in Excel, or compute required values in Python from 01_Input."),
        ],
        [2100, 3500, 3760],
        font_size=8.2,
    )

    add_section_start(doc, "11. Inclusion, Access, and Participation")
    add_bullets(doc, [
        "Allow teams to combine roles based on size, but rotate who reports during class so technical labor is not isolated in one student.",
        "Provide formulas and chart descriptions in text, not color alone. Require labels and patterns sufficient for grayscale review.",
        "Offer a paper-based verification pathway if a student device fails; finance learning should continue even when installation does not.",
        "Share the student assignment before class so students can preview vocabulary and code architecture.",
        "Use structured turn-taking in group reports: data, finance concept, code control, and limitation should be voiced by different members.",
        "Evaluate understanding through explanation and traceability, not typing speed or prior programming experience.",
    ])

    add_section_start(doc, "12. Presentation and Debrief Protocol")
    add_bullets(doc, [
        "Limit each team to five minutes and five content slides.",
        "Require one thesis, one evidence point from each visual, one AI correction, and one caveat.",
        "After each presentation, randomly choose a member for one code/formula question and another for one interpretation question.",
        "Peer listeners record one verified strength and one question about evidence or limitation.",
        "Close by comparing how teams framed the same audited facts differently without changing the numbers.",
    ])
    add_callout(doc, "Debrief question", "What did the AI make faster, and what responsibility did it not remove?", fill=GOLD, accent="7A5A00")

    add_section_start(doc, "Appendix A - Slide Map")
    add_table(
        doc,
        ["Slides", "Section", "Instructor emphasis"],
        [
            ("1-4", "Opening", "Question, outcomes, dashboard, four lenses."),
            ("5-8", "Assignment setup", "Deliverables, roles, source evidence, story workflow."),
            ("9-12", "Sankey", "Income statement identities, visual meaning, audit activity."),
            ("13-18", "FCF", "Profit vs cash, definitions, formula bridge, residual, sign error."),
            ("19-22", "Trend", "Five-year view, CAGR intervals, dual-axis interpretation."),
            ("23-26", "DuPont", "Three factors, averages, leverage, audit activity."),
            ("27-30", "AI prompt and REM", "Junior analyst framing, prompt controls, auditable documentation."),
            ("31-34", "Annotated code", "Architecture, inputs, signs, calculations, stop conditions."),
            ("35-37", "Verification", "Five layers, six checks, documented human correction."),
            ("38-41", "Synthesis and assessment", "Executive story, deliverables, presentation, rubric."),
            ("42-43", "Work session and close", "Team checkpoints and exit ticket."),
        ],
        [1200, 3100, 5060],
        font_size=8.7,
    )

    add_section_start(doc, "Appendix B - Verified Benchmarks")
    add_table(
        doc,
        ["Benchmark", "Expected value"],
        [
            ("FY2025 net sales", "$84,284M"),
            ("FY2025 operating income (EBIT)", "$20,451M"),
            ("FY2025 net earnings", "$16,065M"),
            ("FY2025 cash flow from operations", "$17,817M"),
            ("FY2025 capital expenditures", "$3,773M magnitude"),
            ("Textbook FCF", "$14,866.23M"),
            ("Cash FCF", "$14,044M"),
            ("P&G adjusted FCF", "$14,606M, non-GAAP"),
            ("Sales CAGR", "2.58%"),
            ("GAAP diluted EPS CAGR", "4.30%"),
            ("Simplified DuPont ROE", "31.24%"),
        ],
        [5200, 4160],
        first_col_bold=True,
        font_size=9,
    )
    add_image(
        doc,
        ROOT / "PG_Financial_Story_Dashboard.png",
        6.25,
        "Verified instructor dashboard. Sources: P&G FY2025 and FY2022 Forms 10-K; figures in the supplied package.",
    )

    path = OUT_DIR / "PG_Financial_Story_Instructor_Lesson_Plan.docx"
    doc.save(path)
    return path


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    student = build_student_assignment()
    instructor = build_instructor_guide()
    print(student)
    print(instructor)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
